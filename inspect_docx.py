from docx import Document
import os

path = r"C:\EBA Clearing\APIs\OAS Comparison\RT1 API Participant\2026Q4\OAS_Comparison_Interface_Compatibility_20260324_210401.docx"

if not os.path.exists(path):
    print("Error: File not found.")
    exit(1)

doc = Document(path)

print("--- Document Info ---")
print(f"Number of paragraphs: {len(doc.paragraphs)}")
print(f"Number of tables: {len(doc.tables)}")

print("\n--- First 5 Paragraphs ---")
for i, p in enumerate(doc.paragraphs[:5]):
    print(f"P{i}: {p.text}")

if doc.tables:
    print("\n--- Table 1 (Metadata) ---")
    t1 = doc.tables[0]
    for r in t1.rows:
        print([c.text.strip() for c in r.cells])

    if len(doc.tables) > 1:
        print("\n--- Table 2 (Summary) ---")
        t2 = doc.tables[1]
        print(f"Columns: {len(t2.columns)}")
        # Print first 5 rows
        for i, r in enumerate(t2.rows[:5]):
            print([c.text.strip() for c in r.cells])

    # Find the big discrepancy details table (usually after Summary)
    # Let's check Table 3
    if len(doc.tables) > 2:
        print("\n--- Table 3 (Details/Sample) ---")
        t3 = doc.tables[2]
        print(f"Columns: {len(t3.columns)}")
        for i, r in enumerate(t3.rows[:2]):
            print([c.text.strip() for c in r.cells])
