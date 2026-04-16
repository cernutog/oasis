from pathlib import Path
import shutil

from docx import Document

from src.oas_diff.compatibility_analyzer import CompatibilityIssue
from src.oas_diff.generators.compatibility_generator import CompatibilityDocxGenerator


def _table_text(table) -> str:
    return "\n".join(cell.text for row in table.rows for cell in row.cells)


def test_summary_excludes_added_removed_endpoints():
    issues = [
        CompatibilityIssue(
            path="/removed",
            method="GET",
            location="Endpoint",
            item_name="/removed",
            issue_type="Removed",
            details="Endpoint 'GET /removed' was removed from the new spec.",
        ),
        CompatibilityIssue(
            path="/added",
            method="POST",
            location="Endpoint",
            item_name="/added",
            issue_type="Added",
            details="Endpoint 'POST /added' was added in the new spec.",
        ),
        CompatibilityIssue(
            path="/shared",
            method="GET",
            location="Response 200",
            item_name="requestId",
            issue_type="Constraint Mismatch",
            details="Property 'minLength' changed from 0 to 1",
            old_value=0,
            new_value=1,
        ),
    ]

    temp_dir = Path("tmp/test_compatibility_generator_summary")
    if temp_dir.exists():
        shutil.rmtree(temp_dir)
    temp_dir.mkdir(parents=True, exist_ok=True)

    output_path = temp_dir / "compatibility.docx"
    try:
        CompatibilityDocxGenerator(issues, "old.yaml", "new.yaml").generate(str(output_path))

        doc = Document(output_path)
        assert doc.tables, "Expected at least one table in generated report"

        summary_text = _table_text(doc.tables[1])
        assert "Constraint Mismatch" in summary_text
        assert "Property 'minLength' changed from 0 to 1" in summary_text
        assert "Endpoint 'GET /removed' was removed from the new spec." not in summary_text
        assert "Endpoint 'POST /added' was added in the new spec." not in summary_text

        full_text = "\n".join(par.text for par in doc.paragraphs)
        assert "Added Endpoints" in full_text
        assert "Removed Endpoints" in full_text
        assert "POST /added" in full_text
        assert "GET /removed" in full_text
    finally:
        if temp_dir.exists():
            shutil.rmtree(temp_dir)
