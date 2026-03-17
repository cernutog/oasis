from docx import Document
import os

src_path = r"src\resources\diff_templates\template_impact.docx"
dst_path = r"src\resources\diff_templates\template_compatibility.docx"

if not os.path.exists(src_path):
    print(f"Error: Source template not found at {src_path}")
    exit(1)

doc = Document(src_path)

print("Processing paragraphs...")
for para in doc.paragraphs:
    if "Impact" in para.text:
        print(f"Found in paragraph: {para.text}")
        para.text = para.text.replace("Impact Analysis", "Interface Compatibility")
        para.text = para.text.replace("Impact Report", "Interface Compatibility Report")
        para.text = para.text.replace("Impact", "Interface Compatibility")

print("Processing tables...")
for table in doc.tables:
    for row in table.rows:
        for cell in row.cells:
            for para in cell.paragraphs:
                if "Impact" in para.text:
                    print(f"Found in table: {para.text}")
                    para.text = para.text.replace("Impact Analysis", "Interface Compatibility")
                    para.text = para.text.replace("Impact Report", "Interface Compatibility Report")
                    para.text = para.text.replace("Impact", "Interface Compatibility")

# Process headers/footers
print("Processing headers/footers...")
for section in doc.sections:
    header = section.header
    for para in header.paragraphs:
        if "Impact" in para.text:
             print(f"Found in header: {para.text}")
             para.text = para.text.replace("Impact Report", "Interface Compatibility Report")

doc.save(dst_path)
print(f"Saved template to {dst_path}")
