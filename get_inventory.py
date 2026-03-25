from docx import Document
import os, json

golden_path = r"C:\EBA Clearing\APIs\OAS Comparison\RT1 API Participant\2026Q4\OAS_Comparison_Interface_Compatibility_20260324_210401.docx"
doc = Document(golden_path)

inventory = []
current_endpoint = ""
for p in doc.paragraphs:
    if p.style.name == 'Titolo 2' and ('/' in p.text or ' ' in p.text):
        current_endpoint = p.text

# Iterating Tables to get full list
for table in doc.tables:
    if len(table.columns) == 4 and "Location/Scope" in table.rows[0].cells[0].text:
        for row in table.rows[1:]:
            loc = row.cells[0].text.strip()
            item = row.cells[1].text.strip()
            type_is = row.cells[2].text.strip()
            details = row.cells[3].text.strip()
            # Normalize details (e.g. from To)
            inventory.append({
                "endpoint": current_endpoint,
                "location": loc,
                "item_name": item,
                "type": type_is,
                "details": details
            })

with open("golden_inventory.json", "w", encoding='utf-8') as f:
    json.dump(inventory, f, indent=4)

print(f"INVENTORY COMPLETE: {len(inventory)} rows")
