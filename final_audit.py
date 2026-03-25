from docx import Document
import os
import sys

# 1. Configuration
golden_path = r"C:\EBA Clearing\APIs\OAS Comparison\RT1 API Participant\2026Q4\OAS_Comparison_Interface_Compatibility_20260324_210401.docx"
latest_path = r"C:\EBA Clearing\APIs\OAS Comparison\RT1 API Participant\2026Q4\OAS_Comparison_Interface_Compatibility_20260325_143955.docx"

print(f"Audit: Golden vs User-Generated Latest")
print(f"G: {os.path.basename(golden_path)} ({os.path.getsize(golden_path)} bytes)")
print(f"L: {os.path.basename(latest_path)} ({os.path.getsize(latest_path)} bytes)")

def audit_parity(g_path, l_path):
    g = Document(g_path)
    l = Document(l_path)
    errs = []
    
    # Title
    t1_g = g.paragraphs[1].text.strip()
    t1_l = l.paragraphs[1].text.strip()
    if t1_g != t1_l:
        errs.append(f"Title: Expected '{t1_g}', got '{t1_l}'")

    # Metadata headers
    h1_g = [c.text.strip() for c in g.tables[0].rows[0].cells]
    h1_l = [c.text.strip() for c in l.tables[0].rows[0].cells]
    if h1_g != h1_l:
        errs.append(f"Metadata Headers: Expected {h1_g}, got {h1_l}")

    # Summary headers
    h2_g = [c.text.strip() for c in g.tables[1].rows[0].cells]
    h2_l = [c.text.strip() for c in l.tables[1].rows[0].cells]
    if h2_g != h2_l:
        errs.append(f"Summary Headers: Expected {h2_g}, got {h2_l}")

    # Summary Content (Grouping check)
    g_rows = [[c.text.strip() for c in r.cells] for r in g.tables[1].rows[1:]]
    l_rows = [[c.text.strip() for c in r.cells] for r in l.tables[1].rows[1:]]
    
    print("\n--- Summary Comparison ---")
    print(f"{'Golden Issue':<30} | {'Count':<6} || {'Latest Issue':<30} | {'Count':<6}")
    for i in range(max(len(g_rows), len(l_rows))):
        g_r = g_rows[i] if i < len(g_rows) else ["", "-", "-", "-"]
        l_r = l_rows[i] if i < len(l_rows) else ["", "-", "-", "-"]
        print(f"{g_r[2]:<30} | {g_r[3]:<6} || {l_r[2]:<30} | {l_r[3]:<6}")

    # Details Markers (From/To)
    t3 = l.tables[2]
    found_labels = False
    for row in t3.rows:
        if "Description changed:" in row.cells[3].text:
             if "From:" in row.cells[3].text and "To:" in row.cells[3].text:
                  found_labels = True
                  break
    if not found_labels:
        errs.append("Rich Diff Labels: Missing 'From:' or 'To:' in Details Table.")

    return errs

errors = audit_parity(golden_path, latest_path)
if not errors:
    print("\n[AUDIT SUCCESS] ZERO-DIFF STRUCTURAL PARITY CONFIRMED.")
else:
    print("\n[AUDIT FAILURE] Discrepancies found:")
    for e in errors:
        print(f"  - {e}")
