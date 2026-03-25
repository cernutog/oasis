from docx import Document
import os, json

golden_path = r"C:\EBA Clearing\APIs\OAS Comparison\RT1 API Participant\2026Q4\OAS_Comparison_Interface_Compatibility_20260324_210401.docx"
doc = Document(golden_path)

inventory = []
current_endpoint = ""
for p in doc.paragraphs:
    if p.style.name == 'Titolo 2' and ('/' in p.text or ' ' in p.text):
        current_endpoint = p.text

def map_issue_type(text):
    text = text.lower()
    if "description" in text: return "Description Change"
    if "required" in text: return "Constraint Mismatch"
    if "pattern" in text: return "Constraint Mismatch"
    if "minlength" in text: return "Constraint Mismatch"
    if "maxlength" in text: return "Constraint Mismatch"
    if "enum" in text: return "Enum values order changed"
    return "Constraint Mismatch"

# Iterating Tables to get full list
for table in doc.tables:
    if len(table.columns) == 4 and "Location/Scope" in table.rows[0].cells[0].text:
        for row in table.rows[1:]:
            loc = row.cells[0].text.strip()
            item = row.cells[1].text.strip()
            # type_is = row.cells[2].text.strip() # The type column might be multi-line too
            details_cell = row.cells[3].text.strip()
            
            # Split details by line to capture multi-issue rows
            lines = [l.strip() for l in details_cell.split('\n') if l.strip()]
            for line in lines:
                inventory.append({
                    "endpoint": current_endpoint,
                    "location": loc,
                    "item_name": item,
                    "details": line,
                    "type": map_issue_type(line)
                })

with open("golden_inventory.json", "w", encoding='utf-8') as f:
    json.dump(inventory, f, indent=4)

print(f"DEEP INVENTORY COMPLETE: {len(inventory)} issues found in 243 rows")
