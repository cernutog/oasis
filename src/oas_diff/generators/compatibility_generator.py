from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os
import datetime
import difflib
import re
from typing import List
from ..compatibility_analyzer import CompatibilityIssue

# Helper for OXML
# --- OXML Helpers (Safe Insertion) ---
def get_or_add_child(parent, tag_name, order_list=None):
    if order_list is None:
        order_list = []
    child = parent.find(qn(tag_name))
    if child is not None:
        return child
    child = OxmlElement(tag_name)
    
    if order_list:
        try:
            my_idx = order_list.index(tag_name)
            for i in range(my_idx + 1, len(order_list)):
                next_tag = order_list[i]
                next_element = parent.find(qn(next_tag))
                if next_element is not None:
                    parent.insert(parent.index(next_element), child)
                    return child
        except ValueError:
            pass
            
    parent.append(child)
    return child

TBL_PR_ORDER = ['w:tblStyle', 'w:tblpPr', 'w:tblOverlap', 'w:bidiVisual', 'w:tblStyleRowBandSize', 'w:tblStyleColBandSize', 'w:tblW', 'w:jc', 'w:tblCellSpacing', 'w:tblInd', 'w:tblBorders', 'w:shd', 'w:tblLayout', 'w:tblCellMar', 'w:tblLook']
TC_PR_ORDER = ['w:cnfStyle', 'w:tcW', 'w:gridSpan', 'w:hMerge', 'w:vMerge', 'w:tcBorders', 'w:shd', 'w:noWrap', 'w:tcMar', 'w:textDirection', 'w:tcFitText', 'w:vAlign', 'w:hideMark', 'w:headers', 'w:cellIns', 'w:cellDel', 'w:cellMerge', 'w:tcPrChange']
R_PR_ORDER = ['w:rStyle', 'w:rFonts', 'w:b', 'w:bCs', 'w:i', 'w:iCs', 'w:caps', 'w:smallCaps', 'w:strike', 'w:dstrike', 'w:outline', 'w:shadow', 'w:emboss', 'w:imprint', 'w:noProof', 'w:snapToGrid', 'w:vanish', 'w:webHidden', 'w:color', 'w:spacing', 'w:w', 'w:kern', 'w:position', 'w:sz', 'w:szCs', 'w:highlight', 'w:u', 'w:effect', 'w:bdr', 'w:shd', 'w:fitText', 'w:vertAlign', 'w:rtl', 'w:cs', 'w:em', 'w:lang', 'w:eastAsianLayout', 'w:specVanish', 'w:oMath']

class CompatibilityDocxGenerator:
    def _set_repeat_header(self, row):
        trPr = row._tr.get_or_add_trPr()
        tblHeader = get_or_add_child(trPr, 'w:tblHeader', [])
        return row

    """
    Generates a Word Document (.docx) detailing Interface Compatibility issues.
    """
    def __init__(self, issues: List[CompatibilityIssue], old_path: str, new_path: str, template_path: str = None, spec1: dict = None, spec2: dict = None):
        self.issues = issues
        self.old_path = old_path
        self.new_path = new_path
        self.spec1 = spec1 or {}
        self.spec2 = spec2 or {}
        
        if template_path:
            template_path = template_path.strip('"\' ')
            
        if template_path and os.path.exists(template_path):
            self.doc = Document(template_path)
        else:
            # Fallback to resources directory
            res_dir = os.path.abspath(os.path.join(os.path.dirname(__file__), "..", "..", "resources", "diff_templates"))

            comp_temp = os.path.join(res_dir, "template_compatibility.docx")
            if os.path.exists(comp_temp):
                self.doc = Document(comp_temp)
            elif os.path.exists("template.docx"):
                self.doc = Document("template.docx")
            else:
                self.doc = Document()
            
        self._setup_styles()

    def _setup_styles(self):
        # Normal text
        normal = self.doc.styles['Normal']
        normal.font.name = 'Segoe UI'
        normal.font.size = Pt(10)

        # Title
        if 'Title' in self.doc.styles:
            style = self.doc.styles['Title']
            style.font.name = 'Georgia'
            style.font.size = Pt(22)
            style.font.bold = True
            style.font.color.rgb = RGBColor(31, 78, 121)

    def generate(self, output_path: str):
        """Builds the document structure and saves it."""
        self.doc.add_heading('OpenAPI Comparison - Interface Changelog Report', 0)


        
        # 1. Metadata Table
        self._add_metadata_table()

        # 2. Executive Summary or Total Count

        
        if not self.issues:
             self.doc.add_paragraph("No interface compatibility discrepancies were found. The specifications are fully compatible in terms of request/response payloads and parameters.")
             self.doc.save(output_path)
             return

        # 3. List of Issues (Table grouped by Path/Method)

        # 4. Summary Section

        self.doc.add_heading('Summary', 1)
        
        # Summary logic: aggregate by (type, details, old_value, new_value)
        freq = {}
        # Calculate frequency and collect for summary
        freq = {}
        for issue in self.issues:
            # Clean representation for lists (Enums)
            if isinstance(issue.old_value, list):
                v1 = ", ".join(map(str, issue.old_value))
            else:
                v1 = str(issue.old_value) if issue.old_value is not None else "<None>"
                
            if isinstance(issue.new_value, list):
                v2 = ", ".join(map(str, issue.new_value))
            else:
                v2 = str(issue.new_value) if issue.new_value is not None else "<None>"
            
            key = (issue.issue_type, issue.details, v1, v2)
            freq[key] = freq.get(key, 0) + 1
             
        self.doc.add_paragraph().add_run(f'Total Issues Found: {len(self.issues)}').bold = True
             
        sorted_freq = sorted(freq.items(), key=lambda x: x[1], reverse=True)
        widths_s = [Inches(0.4), Inches(2.1), Inches(3.5), Inches(1.0)]
        table = self._create_table(4, widths_s)
        self._set_repeat_header(table.rows[0])
        
        # Style header row
        for i, header_text in enumerate(["#", "Issue Type", "Details (Variation)", "Count"]):
            t = table.cell(0, i)
            t.text = header_text
            tcPr = t._tc.get_or_add_tcPr()
            shd = get_or_add_child(tcPr, 'w:shd')
            shd.set(qn('w:val'), 'clear')
            shd.set(qn('w:fill'), '1F4E79')
            if t.paragraphs and t.paragraphs[0].runs:
                t.paragraphs[0].runs[0].font.bold = True
                t.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
            
        issue_id_map = {}
        for idx, (key_tuple, count) in enumerate(sorted_freq):
             issue_type, details, v1, v2 = key_tuple
             issue_id_map[key_tuple] = idx + 1
             row = table.add_row().cells
             row[0].text = f"[{idx + 1}]"
             row[1].text = issue_type
             
             cell_var = row[2]
             p_var = cell_var.paragraphs[0]
             
             if issue_type in ["Description Change", "Enum values order changed"]:
                  p_var.add_run(f"{details}.\n")
                  p_var.add_run("Old: ").bold = True
                  # We use a second paragraph for New to keep it clean
                  p_new = cell_var.add_paragraph()
                  p_new.add_run("New: ").bold = True
                  self._render_rich_diff(p_var, p_new, v1, v2)
             else:
                  # Use semantic shading for constraints in summary too
                  # msg example: "Constraint 'pattern' changed from 'abc' to <None>"
                  match = re.search(r"((?:Constraint|Property) '.*?' changed from )(.*?) ( to )(.*)$", details)
                  if match:
                       p_var.add_run(match.group(1))
                       r_v1 = p_var.add_run(match.group(2))
                       if v1 != "<None>" and v2 != "<None>": self._apply_shading(r_v1, "FFF3CD")
                       elif v2 == "<None>": self._apply_shading(r_v1, "F8D7DA")
                       
                       p_var.add_run(match.group(3))
                       
                       r_v2 = p_var.add_run(match.group(4))
                       if v1 != "<None>" and v2 != "<None>": self._apply_shading(r_v2, "FFF3CD")
                       elif v1 == "<None>": self._apply_shading(r_v2, "D4EDDA")
                  else:
                       p_var.add_run(details)


             row[3].text = str(count)
             for j in range(4):
                 self._style_body_cell(row[j])

        self.doc.add_heading('Discrepancy Details', 1)
        
        # Group issues by endpoint for better flow
        grouped_issues = {}
        for issue in self.issues:
             key = f"{issue.method} {issue.path}"
             if key not in grouped_issues:
                  grouped_issues[key] = []
             grouped_issues[key].append(issue)

        def _format_item_name(name: str) -> str:
            if not name: return "-"
            parts = name.split('.')
            res = []
            for i, p in enumerate(parts):
                res.append("  " * i + p)
            return "\n".join(res)

        for endpoint, op_issues in sorted(grouped_issues.items(), key=lambda x: (x[0].split()[1], x[0].split()[0])):
            self.doc.add_heading(endpoint, 2)
            
            # Group by (location, item_name) — keep FULL issue objects
            dedup = {}
            for issue in op_issues:
                k = (issue.location, issue.item_name)
                dedup.setdefault(k, []).append(issue)

            widths = [Inches(1.5), Inches(1.5), Inches(1.2), Inches(2.8)]
            table = self._create_table(4, widths)
            self._add_all_borders(table)
            
            # Headers
            headers = ['Location/Scope', 'Property/Param', 'Issue Type', 'Details']
            h_row = table.rows[0].cells
            self._set_repeat_header(table.rows[0])
            for i, h in enumerate(headers):
                h_row[i].text = h
                tcPr = h_row[i]._tc.get_or_add_tcPr()
                shd = get_or_add_child(tcPr, 'w:shd')
                shd.set(qn('w:val'), 'clear')
                shd.set(qn('w:fill'), '1F4E79')
                p = h_row[i].paragraphs[0]
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                if p.runs:
                    p.runs[0].font.bold = True
                    p.runs[0].font.color.rgb = RGBColor(255, 255, 255)

            # Rows
            for (loc, item), group in dedup.items():
                # Sort group: Description Change first
                group.sort(key=lambda x: 0 if x.issue_type == "Description Change" else 1)
                
                row_cells = table.add_row().cells
                row_cells[0].text = loc
                row_cells[1].text = _format_item_name(item)
                
                # Issue Type column — unique types
                unique_types = list(dict.fromkeys([iss.issue_type for iss in group]))
                it_cell = row_cells[2]
                for i_t, t_name in enumerate(unique_types):
                    p_it = it_cell.paragraphs[0] if i_t == 0 else it_cell.add_paragraph()
                    p_it.text = t_name

                # Details column — each issue gets its own paragraph
                cell = row_cells[3]
                for idx, issue in enumerate(group):
                    v1 = str(issue.old_value) if issue.old_value is not None else "<None>"
                    v2 = str(issue.new_value) if issue.new_value is not None else "<None>"
                    
                    key_tuple = (issue.issue_type, issue.details, v1, v2)
                    s_idx = issue_id_map.get(key_tuple, "")
                    ref_str = f"[{s_idx}]" if s_idx else ""
                    
                    p_desc = cell.paragraphs[0] if idx == 0 and not cell.paragraphs[0].text else cell.add_paragraph()
                    p_desc.paragraph_format.left_indent = Inches(0.22)
                    p_desc.paragraph_format.first_line_indent = Inches(-0.22)
                    
                    # Highlight simple constraints within the text if possible
                    if issue.issue_type != "Description Change" and (issue.old_value is not None or issue.new_value is not None):
                         p_desc.add_run(f"{ref_str} ")
                         msg = issue.details
                         # Pattern matches: Constraint/Property 'xxx' changed from 'yyy' to 'zzz'
                         match = re.search(r"((?:Constraint|Property) '.*?' changed from )(.*?) ( to )(.*)$", msg)
                         if match:
                              p_desc.add_run(match.group(1))
                              r_v1 = p_desc.add_run(match.group(2))
                              if v1 is not None and v2 is not None: self._apply_shading(r_v1, "FFF3CD") # Yellow (mod)
                              elif v2 is None: self._apply_shading(r_v1, "F8D7DA") # Red (removed)
                              
                              p_desc.add_run(match.group(3))
                              
                              r_v2 = p_desc.add_run(match.group(4))
                              if v1 is not None and v2 is not None: self._apply_shading(r_v2, "FFF3CD") # Yellow (mod)
                              elif v1 is None: self._apply_shading(r_v2, "D4EDDA") # Green (added)
                         else:
                              p_desc.add_run(msg)
                    else:
                         p_desc.add_run(f"{ref_str} {issue.details}")


                    # Old:/New: value lines
                    if issue.issue_type in ["Description Change", "Enum values order changed"]:
                        p_old = cell.add_paragraph()
                        p_old.add_run("Old:").bold = True
                        p_old.add_run(" ")
                        p_new = cell.add_paragraph()
                        p_new.add_run("New:").bold = True
                        p_new.add_run(" ")
                        
                        # Convert lists to strings for diffing (e.g. for enums)
                        val1 = ", ".join(map(str, issue.old_value)) if isinstance(issue.old_value, list) else v1
                        val2 = ", ".join(map(str, issue.new_value)) if isinstance(issue.new_value, list) else v2
                        
                        self._render_rich_diff(p_old, p_new, val1, val2)
                    
                    # Space between issues in same cell
                    if idx < len(group) - 1:
                        cell.add_paragraph()

        self.doc.save(output_path)

    def _add_all_borders(self, table):
        tbl = table._tbl
        tblPr = tbl.tblPr
        tblBorders = get_or_add_child(tblPr, 'w:tblBorders', TBL_PR_ORDER)
        for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
            border = get_or_add_child(tblBorders, f'w:{border_name}')
            border.set(qn('w:val'), 'single')
            border.set(qn('w:sz'), '4')
            border.set(qn('w:color'), '000000')

    def _create_table(self, cols, widths):
        table = self.doc.add_table(rows=1, cols=cols)
        self._remove_all_borders(table)
        
        total_width = sum(w.inches for w in widths)
        self._set_table_fixed_width(table, total_width)
        table.alignment = WD_TABLE_ALIGNMENT.LEFT
        
        tblPr = table._tblPr
        tblGrid = table._element.find(qn('w:tblGrid'))
        if tblGrid is None:
            tblGrid = OxmlElement('w:tblGrid')
            table._element.insert(table._element.index(tblPr) + 1, tblGrid)
        else:
            tblGrid.clear()
            
        for w in widths:
            col = OxmlElement('w:gridCol')
            col.set(qn('w:w'), str(int(w.twips)))
            tblGrid.append(col)
            
        for i, cell in enumerate(table.rows[0].cells):
            cell.width = widths[i]
            
        return table

    def _remove_all_borders(self, table):
        tbl = table._tbl
        tblPr = tbl.tblPr
        tblBorders = get_or_add_child(tblPr, 'w:tblBorders', TBL_PR_ORDER)
        for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
            border = get_or_add_child(tblBorders, f'w:{border_name}')
            border.set(qn('w:val'), 'nil')

    def _set_table_fixed_width(self, table, width_inches):
        tbl = table._tbl
        tblPr = tbl.tblPr
        tblW = get_or_add_child(tblPr, 'w:tblW', TBL_PR_ORDER)
        tblW.set(qn('w:w'), str(int(width_inches * 1440)))
        tblW.set(qn('w:type'), 'dxa')
        tblLayout = get_or_add_child(tblPr, 'w:tblLayout', TBL_PR_ORDER)
        tblLayout.set(qn('w:type'), 'fixed')

    def _style_body_cell(self, cell):
        # Horizontal Border Only (Light Grey)
        tcPr = cell._tc.get_or_add_tcPr()
        tcBorders = get_or_add_child(tcPr, 'w:tcBorders', TC_PR_ORDER)
        for side in ['top', 'left', 'right']:
            tag = get_or_add_child(tcBorders, f'w:{side}')
            tag.set(qn('w:val'), 'nil')
        bottom = get_or_add_child(tcBorders, 'w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), '4')
        bottom.set(qn('w:color'), 'E0E0E0')

    def _add_metadata_table(self):
        widths = [Inches(1.5), Inches(2.75), Inches(2.75)]
        table = self._create_table(3, widths)
        # Custom Header Styling (Dark Blue Background, White Text)
        headers = ['Detail', 'Old Specification', 'New Specification']
        row = table.rows[0]
        
        for i, text in enumerate(headers):
             cell = row.cells[i]
             cell.text = text
             p = cell.paragraphs[0]
             if p.runs:
                 p.runs[0].font.bold = True
                 p.runs[0].font.color.rgb = RGBColor(255, 255, 255) # White
             
             # Dark Blue Background
             tcPr = cell._tc.get_or_add_tcPr()
             shd = get_or_add_child(tcPr, 'w:shd')
             shd.set(qn('w:val'), 'clear')
             shd.set(qn('w:fill'), '1F4E79') # Dark Blue

             # Bottom border for header like analytic report
             tcBorders = get_or_add_child(tcPr, 'w:tcBorders')
             bottom = get_or_add_child(tcBorders, 'w:bottom')
             bottom.set(qn('w:val'), 'single')
             bottom.set(qn('w:sz'), '12')
             bottom.set(qn('w:color'), 'FFFFFF')

        def get_info(spec, path):
             info = spec.get('info', {}) if isinstance(spec, dict) else {}
             return {
                 'file': os.path.basename(path) if path else "N/A",
                 'title': info.get('title', 'N/A'),
                 'version': str(info.get('version', 'N/A'))
             }
             
        old_info = get_info(self.spec1, self.old_path)
        new_info = get_info(self.spec2, self.new_path)

        rows = [
             ("File Name", old_info['file'], new_info['file']),
             ("API Title", old_info['title'], new_info['title']),
             ("Version", old_info['version'], new_info['version'])
        ]
        
        for d, o, n in rows:
             r_cells = table.add_row().cells
             r_cells[0].text = d
             r_cells[1].text = o
             r_cells[2].text = n
             if r_cells[0].paragraphs:
                 r_cells[0].paragraphs[0].runs[0].font.bold = True
             for j in range(3):
                 self._style_body_cell(r_cells[j])

    def _norm_desc(self, txt):
        """Collapses all whitespace into single spaces and strips."""
        if txt is None: return ''
        return re.sub(r'\s+', ' ', str(txt)).strip()

    def _render_rich_diff(self, p_old, p_new, text_old, text_new):
        """Renders word-level diff with semantic coloring (Red for removal, Green for addition)."""
        # 1. Standardize text to eliminate whitespace-only differences
        v1 = self._norm_desc(text_old)
        v2 = self._norm_desc(text_new)
        
        # Handle <None> or null cases as pure addition/deletion
        is_none1 = not v1 or v1 == "<None>"
        is_none2 = not v2 or v2 == "<None>"
        
        if is_none1 and is_none2:
            p_old.add_run("<None>")
            p_new.add_run("<None>")
            return
        if is_none1:
            p_old.add_run("<None>")
            r_n = p_new.add_run(v2)
            self._apply_shading(r_n, 'D4EDDA') # Green
            return
        if is_none2:
            r_o = p_old.add_run(v1)
            self._apply_shading(r_o, 'F8D7DA') # Red
            p_new.add_run("<None>")
            return

        # 2. Word-level diff on normalized text
        words1 = v1.split()
        words2 = v2.split()
        
        # Disable autojunk for precise technical text comparison
        sm = difflib.SequenceMatcher(None, words1, words2, autojunk=False)
        for tag, i1, i2, j1, j2 in sm.get_opcodes():
            if tag == 'equal':
                for word in words1[i1:i2]:
                    p_old.add_run(word + " ")
                    p_new.add_run(word + " ")
            elif tag == 'replace':
                # Deleted parts in red (old)
                r_old = p_old.add_run(" ".join(words1[i1:i2]) + " ")
                self._apply_shading(r_old, "F8D7DA")
                # Added parts in green (new)
                r_new = p_new.add_run(" ".join(words2[j1:j2]) + " ")
                self._apply_shading(r_new, "D4EDDA")
            elif tag == 'delete':
                r_old = p_old.add_run(" ".join(words1[i1:i2]) + " ")
                self._apply_shading(r_old, "F8D7DA")
            elif tag == 'insert':
                r_new = p_new.add_run(" ".join(words2[j1:j2]) + " ")
                self._apply_shading(r_new, "D4EDDA")

    def _apply_shading(self, run, color):
        """Applies background shading color to a run."""
        rPr = run._r.get_or_add_rPr()
        shd = get_or_add_child(rPr, 'w:shd', R_PR_ORDER)
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:fill'), color)
        return run
