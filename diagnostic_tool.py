# Mega Diagnostic: Exact set comparison between Golden and Generated
import sys, os, yaml, re
from docx import Document

sys.path.insert(0, os.getcwd())
from src.oas_diff.resolver import resolve_spec
from src.oas_diff.compatibility_analyzer import CompatibilityAnalyzer, CompatibilityIssue

OLD = r"C:/EBA Clearing/APIs/Source OAS/20260323/EBACL_RT1_20260323_Openapi3.0-SWIFT_RT1_API_Participants_5_0_v20260323.yaml"
NEW = r"C:/EBA Clearing/APIs/Generated OAS/RT1 API Participants/2026Q4/generated_oas_3.0_SWIFT.yaml"
GOLD = r"C:\EBA Clearing\APIs\OAS Comparison\RT1 API Participant\2026Q4\OAS_Comparison_Interface_Compatibility_20260324_210401.docx"

def normalize_prop_path(p):
    # Collapse spaces/newlines into dots to match our dot-notation
    p = p.replace('\r', '')
    lines = [line.strip() for line in p.split('\n') if line.strip()]
    return ".".join(lines)

# 1. Extract from Golden with better path reconstruction
def extract_golden_data(path):
    doc = Document(path)
    records = []
    current_ep = None
    
    from docx.oxml.table import CT_Tbl
    from docx.oxml.text.paragraph import CT_P
    from docx.table import Table
    from docx.text.paragraph import Paragraph

    def iter_blocks(parent):
        for child in parent.element.body:
            if isinstance(child, CT_P): yield Paragraph(child, parent)
            elif isinstance(child, CT_Tbl): yield Table(child, parent)

    for block in iter_blocks(doc):
        if isinstance(block, Paragraph) and block.style.name == 'Heading 2':
            current_ep = block.text.strip()
        elif isinstance(block, Table):
            if block.rows[0].cells[0].text.strip() == 'Location/Scope':
                for r in block.rows[1:]:
                    loc = r.cells[0].text.strip()
                    prop_raw = r.cells[1].text.strip()
                    norm_prop = normalize_prop_path(prop_raw)
                    types = [p.text.strip() for p in r.cells[2].paragraphs if p.text.strip()]
                    details = "\n".join([p.text for p in r.cells[3].paragraphs])
                    
                    for t in types:
                         # Distinguish pattern vs others in Constraint Mismatch
                         actual_type = t
                         if t == 'Constraint Mismatch':
                             if 'pattern' in details: actual_type = 'pattern'
                             elif 'required' in details:
                                 if 'False to True' in details: actual_type = 'req_true'
                                 else: actual_type = 'req_false'
                             elif 'minLength' in details: actual_type = 'minLength'
                             elif 'maxLength' in details: actual_type = 'maxLength'
                         
                         records.append({'ep': current_ep, 'loc': loc, 'prop': norm_prop, 'type': actual_type})
    return records

gold_records = extract_golden_data(GOLD)

# 2. Run Analyzer
with open(OLD, 'r', encoding='utf-8') as f: s1 = yaml.safe_load(f)
with open(NEW, 'r', encoding='utf-8') as f: s2 = yaml.safe_load(f)
r1 = resolve_spec(s1)
r2 = resolve_spec(s2)

# Temporary Analyzer fix for allOf during diagnostic
def _flatten(s):
    if not isinstance(s, dict): return s
    if 'allOf' in s:
        res = {}
        for item in s['allOf']:
            f_item = _flatten(item)
            if isinstance(f_item, dict):
                for k, v in f_item.items():
                    if k == 'properties': res.setdefault('properties', {}).update(v)
                    elif k == 'required': 
                        res_req = res.setdefault('required', [])
                        res_req.extend([x for x in v if x not in res_req])
                    else: res[k] = v
        for k,v in s.items():
            if k != 'allOf':
                if k == 'properties': res.setdefault('properties', {}).update(v)
                else: res[k] = v
        return res
    return s

class MegaAnalyzer(CompatibilityAnalyzer):
    def _compare_schemas(self, path, method, location, s1, s2, prefix, visited=None, skip_description=False):
        s1 = _flatten(s1)
        s2 = _flatten(s2)
        # Re-run normalizing function as expected by parent
        super()._compare_schemas(path, method, location, s1, s2, prefix, visited, skip_description)

analyzer = MegaAnalyzer(r1, r2)
issues = analyzer.analyze()

gen_records = []
for iss in issues:
    t = iss.issue_type
    if "pattern" in iss.details: t = "pattern"
    elif "required" in iss.details:
        if "False to True" in iss.details: t = "req_true"
        else: t = "req_false"
    elif "minLength" in iss.details: t = "minLength"
    elif "maxLength" in iss.details: t = "maxLength"
    
    gen_records.append({'ep': f"{iss.method} {iss.path}", 'loc': iss.location, 'prop': iss.item_name, 'type': t})

# 3. Compare sets
def to_set(recs): return set((r['ep'], r['prop'], r['type']) for r in recs)

gold_set = to_set(gold_records)
gen_set = to_set(gen_records)

missing = gold_set - gen_set
extra = gen_set - gold_set

print(f"Gold: {len(gold_set)} | Gen: {len(gen_set)}")
print(f"Missing: {len(missing)} | Extra: {len(extra)}")

if missing:
    print("\n--- MISSING ---")
    for m in sorted(list(missing))[:15]: print(f"  {m}")
if extra:
    print("\n--- EXTRA ---")
    for e in sorted(list(extra))[:15]: print(f"  {e}")
